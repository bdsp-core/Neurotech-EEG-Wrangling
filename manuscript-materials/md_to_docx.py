#!/usr/bin/env python3
"""Convert manuscript-draft.md to a Word document with embedded figures."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

URL_RE = re.compile(r'(https?://[^\s)]+)')

BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "manuscript-draft.md"
OUT_PATH = BASE / "Neurotech_EEG_Dataset_Draft.docx"

# Map figure numbers (as they appear in the text) to image files
FIGURE_FILES = {
    "1": BASE / "figures" / "figure1_pipeline.png",
    "2": BASE / "figures" / "figure2_dataset_positioning.png",
    "3": BASE / "figures" / "figure3_patient_characteristics.png",
    "4": BASE / "figures" / "figure4_eeg_findings.png",
}

# Supplementary figures — embedded in the Supplementary Material section
SUPP_FIGURE_FILES = {
    # Supp Fig 1 is the example EEG traces (matches the legend); it was previously
    # mis-wired to the stale annotation-summary bar chart. Regenerate with
    # manuscript-materials/make_supp_figure1_eeg.py.
    "1": BASE / "figures" / "supp_figure1_eeg_example.png",
    "2": BASE / "figures" / "supp_figure2_deidentification.png",
    "3": BASE / "figures" / "supp_figure3_comorbidities_meds.png",
    "4": BASE / "figures" / "supp_figure4_monitoring.png",
}

def set_run_font(run, size=11, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_hyperlink(paragraph, url, text, size=11):
    """Add a real, clickable Word hyperlink (blue, underlined) to the paragraph."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_plain_with_links(paragraph, text, size, bold, italic):
    """Emit a plain text run, turning bare URLs into clickable hyperlinks."""
    for seg in URL_RE.split(text):
        if URL_RE.fullmatch(seg):
            m = re.search(r'[.,;:]+$', seg)          # keep trailing sentence punctuation out of the link
            trail = m.group(0) if m else ''
            url = seg[:len(seg) - len(trail)] if trail else seg
            add_hyperlink(paragraph, url, url, size=size)
            if trail:
                set_run_font(paragraph.add_run(trail), size=size, bold=bold, italic=italic)
        elif seg:
            set_run_font(paragraph.add_run(seg), size=size, bold=bold, italic=italic)


def add_formatted_text(paragraph, text, size=11, bold=False, italic=False):
    """Add text to a paragraph, handling markdown bold/italic and clickable URLs."""
    # Process bold (**text**) and italic (*text*) patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, bold=bold, italic=True)
        else:
            _add_plain_with_links(paragraph, part, size, bold, italic)

def parse_table(lines):
    """Parse markdown table lines into list of rows (each row is list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator rows
        if re.match(r'^\|[\s\-|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows

def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                # Header row in bold
                is_bold = (i == 0) or cell_text.startswith('**')
                clean_text = cell_text.replace('**', '')
                # Handle backticks
                clean_text = clean_text.replace('`', '')
                run = p.add_run(clean_text)
                set_run_font(run, size=9, bold=is_bold)
    return table

def add_figure(doc, fig_num, supplementary=False):
    """Add a figure image if available."""
    source = SUPP_FIGURE_FILES if supplementary else FIGURE_FILES
    fig_path = source.get(str(fig_num))
    if fig_path and fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_path), width=Inches(6.0))
        return True
    return False

def process_markdown(md_text):
    """Parse markdown and generate Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    lines = md_text.split('\n')
    i = 0

    # Track which figures we've inserted
    figures_inserted = set()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            # Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = stripped[2:]
            run = p.add_run(text)
            set_run_font(run, size=16, bold=True)
            i += 1
            continue

        if stripped.startswith('## '):
            text = stripped[3:]
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                set_run_font(run, size=14, bold=True)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
            i += 1
            continue

        # Table (collect all table lines)
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            doc.add_paragraph()  # spacing after table
            continue

        # Numbered list items — preserve the EXPLICIT source number rather than using
        # Word's 'List Number' style, whose auto-numbering runs continuously across the
        # whole document (that made independent lists and the reference list mis-number,
        # e.g. references starting at 8 instead of 1).
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)  # hanging indent
            add_formatted_text(p, stripped)   # keep the "N. " prefix verbatim
            i += 1
            continue

        # Bullet list items
        if stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Handle backtick code in list items
            text = text.replace('`', '')
            add_formatted_text(p, text)
            i += 1
            continue

        # Figure legend lines - detect and potentially insert figure image before legend
        fig_legend_match = re.match(r'\*\*Figure (\d+)\.\*\*', stripped)
        if fig_legend_match:
            fig_num = fig_legend_match.group(1)
            if fig_num not in figures_inserted:
                if add_figure(doc, fig_num):
                    figures_inserted.add(fig_num)

            # Add the legend text
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Supplementary figure legends - insert image then legend
        supp_fig_match = re.match(r'\*\*Supplementary Figure (\d+)\.\*\*', stripped)
        if supp_fig_match:
            supp_num = supp_fig_match.group(1)
            add_figure(doc, supp_num, supplementary=True)
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Table caption lines (bold text starting with Table or Supplementary Table)
        if stripped.startswith('**Table') or stripped.startswith('**Supplementary Table'):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Regular paragraph - handle superscript notation, backticks, etc.
        text = stripped
        # Clean up pandoc-style superscripts (^text^)
        text = re.sub(r'\^(\d+)\^', r'[\1]', text)  # Convert ^1^ to [1]
        text = re.sub(r'\^(\d+-\d+)\^', r'[\1]', text)  # Convert ^1-3^ to [1-3]
        # Clean up backticks
        text = text.replace('`', '')

        p = doc.add_paragraph()

        # Handle italic lines (like word count)
        if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
            run = p.add_run(text[1:-1])
            set_run_font(run, italic=True)
        else:
            add_formatted_text(p, text)

        i += 1

    return doc


def main():
    md_text = MD_PATH.read_text()
    doc = process_markdown(md_text)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.save(str(OUT_PATH))
    print(f"Word document saved to {OUT_PATH}")
    print(f"File size: {OUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
