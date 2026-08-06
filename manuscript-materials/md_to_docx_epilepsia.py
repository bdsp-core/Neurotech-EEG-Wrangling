#!/usr/bin/env python3
"""Render manuscript-epilepsia.md to an Epilepsia-submission-shaped Word document:
title page, double-spacing, continuous line numbers, 12 pt Times New Roman.

Reuses the shared inline renderers (bold/italic, ^superscript^ citations, hyperlinks,
tables, embedded figures) from md_to_docx.py.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE))
import md_to_docx as m   # shared helpers: set_run_font, add_formatted_text, parse_table, add_table_to_doc, add_figure

MD = BASE / "manuscript-epilepsia.md"
OUT = BASE / "Neurotech_EEG_Dataset_Epilepsia.docx"


def set_double_spacing(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)


def add_line_numbers(doc):
    sectPr = doc.sections[0]._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sectPr.append(ln)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    run._r.append(br)


def render_title_page(doc, block):
    """block: the lines between <!--TITLEPAGE--> and <!--ENDTITLEPAGE-->."""
    for raw in block.split("\n"):
        s = raw.strip()
        if not s or s.startswith("<!--"):
            continue
        if s.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            m.set_run_font(p.add_run(s[2:]), size=15, bold=True)
        else:
            p = doc.add_paragraph()
            m.add_formatted_text(p, s, size=12)
    page_break(doc)


def render_body(doc, body):
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]; s = line.strip()
        if not s or s == "---":
            i += 1; continue
        # section heading (## ) and subsection (### )
        if s.startswith("### "):
            p = doc.add_paragraph(); m.set_run_font(p.add_run(s[4:]), size=12, bold=True, italic=True)
            i += 1; continue
        if s.startswith("## "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
            m.set_run_font(p.add_run(s[3:]), size=13, bold=True)
            i += 1; continue
        # table
        if s.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            rows = m.parse_table(tbl)
            m.add_table_to_doc(doc, rows)
            doc.add_paragraph()
            continue
        # figure legend -> insert figure image then legend
        fig = re.match(r"\*\*Figure (\d+)\.\*\*", s)
        if fig:
            if m.add_figure(doc, fig.group(1)):
                pass
            p = doc.add_paragraph(); m.add_formatted_text(p, s, size=12)
            i += 1; continue
        # numbered list (references)
        if re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)
            m.add_formatted_text(p, s, size=12)
            i += 1; continue
        # bullet list (Key Points)
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            m.add_formatted_text(p, s[2:], size=12)
            i += 1; continue
        # normal paragraph (may be a **Table N.** caption or prose)
        p = doc.add_paragraph()
        m.add_formatted_text(p, s, size=12)
        i += 1


def main():
    text = MD.read_text()
    # fill the main-text word count (Introduction..Conclusions)
    body_for_count = text
    a = body_for_count.find("## 1. Introduction")
    b = body_for_count.find("## Author Contributions")
    wc = len(re.findall(r"\S+", re.sub(r"[#*`|>_-]", " ", body_for_count[a:b]))) if a >= 0 and b >= 0 else 0
    text = text.replace("WORDCOUNT_PLACEHOLDER", f"~{wc:,}")

    # split title page from body
    tp = re.search(r"<!--TITLEPAGE-->(.*?)<!--ENDTITLEPAGE-->", text, re.S)
    title_block = tp.group(1) if tp else ""
    body = text[tp.end():] if tp else text

    doc = Document()
    set_double_spacing(doc)
    render_title_page(doc, title_block)
    render_body(doc, body)
    add_line_numbers(doc)
    doc.save(OUT)
    print(f"saved {OUT}  (main-text word count ~{wc:,})")


if __name__ == "__main__":
    main()
